"""Generate Google Ads API Design Document as .docx"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# -- Styles --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

# -- Title --
title = doc.add_heading('Google Ads API — Design Document', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('Piano Tuning KC — Basic Access Application')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph('')  # spacer

# -- Company Info --
doc.add_heading('Company Information', level=1)
info_table = doc.add_table(rows=3, cols=2, style='Light Grid Accent 1')
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cells = [
    ('Company Name', 'Piano Tuning KC'),
    ('Company Website', 'https://pianotuningkc.com'),
    ('Google Ads Customer ID', '790-110-2348'),
]
for i, (label, value) in enumerate(cells):
    info_table.rows[i].cells[0].text = label
    info_table.rows[i].cells[1].text = value
    for cell in info_table.rows[i].cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(2)

doc.add_paragraph('')

# -- Business Model --
doc.add_heading('Business Model', level=1)
doc.add_paragraph(
    'Piano Tuning KC is a local piano tuning service operating in the Kansas City '
    'metropolitan area. We operate a single website (https://pianotuningkc.com) which '
    'serves as our primary online presence for booking piano tuning appointments. '
    'We only advertise for our own business and do not manage ads for any other '
    'companies or clients.'
)

# -- Tool Access / Use --
doc.add_heading('Tool Access / Use', level=1)
doc.add_paragraph(
    'Our tool is an internal-only reporting and automation system used exclusively '
    'by the business owner and a technical consultant to:'
)

uses = [
    ('View and analyze ad performance',
     'A lightweight internal dashboard that displays campaign metrics (clicks, '
     'impressions, conversions, cost, CPC, CTR) across configurable date ranges. '
     'This dashboard is accessible only to authorized internal users and is not '
     'publicly accessible.'),
    ('Generate performance reports',
     'The tool can export ad performance summaries for internal review. These reports '
     'are used solely for internal business decisions (budget allocation, keyword '
     'optimization, seasonal planning). Reports are not shared with any external '
     'agencies or third parties.'),
    ('Automated campaign monitoring',
     'A scheduled script runs periodically to pull the latest campaign metrics and '
     'flag significant changes in performance (e.g., cost-per-conversion exceeding '
     'a threshold, conversion rate drops, budget pacing issues). Alerts are sent to '
     'the business owner via internal notification.'),
]
for title_text, desc in uses:
    p = doc.add_paragraph()
    run = p.add_run(f'{title_text} — ')
    run.bold = True
    p.add_run(desc)
    p.style = doc.styles['List Bullet']

p = doc.add_paragraph()
run = p.add_run('No external parties will have access to the tool.')
run.bold = True
p.add_run(' It is hosted locally and used only by internal personnel.')

# -- Tool Design --
doc.add_heading('Tool Design', level=1)

doc.add_heading('Data Flow Architecture', level=2)
# Create a simple text-based flow diagram
flow = doc.add_paragraph()
flow.alignment = WD_ALIGN_PARAGRAPH.CENTER
flow_text = (
    'Google Ads API (GAQL queries)\n'
    '          ↓\n'
    'Python Script (google-ads library)\n'
    '          ↓\n'
    'Local SQLite Database (metrics cache)\n'
    '          ↓\n'
    'Internal Dashboard (HTML/JS, localhost)\n'
    '          ↓\n'
    'Business Owner views reports'
)
run = flow.add_run(flow_text)
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_heading('How It Works', level=2)

steps = [
    ('Data Collection',
     'A Python script using the official google-ads Python client library makes '
     'read-only API calls using Google Ads Query Language (GAQL) to retrieve campaign, '
     'ad group, keyword, and search term performance data.'),
    ('Data Storage',
     'Retrieved metrics are stored in a local SQLite database to enable historical '
     'trend analysis and reduce redundant API calls. The database is stored locally '
     'and is not accessible externally.'),
    ('Dashboard Display',
     'A simple internal web dashboard (served on localhost) reads from the local '
     'database to display account-level summary metrics, campaign-level performance '
     'comparisons, keyword performance data, and date range filtering.'),
    ('Automated Monitoring',
     'A scheduled script runs every 4 hours during business hours to pull the latest '
     'performance metrics, compare against historical baselines, and log any significant '
     'performance changes.'),
]
for i, (step_title, step_desc) in enumerate(steps, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. {step_title}: ')
    run.bold = True
    p.add_run(step_desc)

doc.add_heading('Data Retention', level=2)
retention_items = [
    'Metrics data is retained locally for up to 12 months for trend analysis.',
    'No personally identifiable information (PII) is collected or stored.',
    'Only aggregate performance metrics are retrieved from the API.',
]
for item in retention_items:
    doc.add_paragraph(item, style='List Bullet')

# -- API Services --
doc.add_heading('API Services Called', level=1)

p = doc.add_paragraph()
run = p.add_run('We will use read-only access ')
run.bold = True
p.add_run('to the following API resources:')

api_table = doc.add_table(rows=9, cols=3, style='Light Grid Accent 1')
api_table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['API Resource', 'Purpose', 'Operations']
for i, h in enumerate(headers):
    api_table.rows[0].cells[i].text = h
    for paragraph in api_table.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

api_rows = [
    ('GoogleAdsService', 'Execute GAQL queries to pull performance reports', 'Search, SearchStream'),
    ('Customer', 'Retrieve account-level settings and summary metrics', 'Read only'),
    ('Campaign', 'Pull campaign names, statuses, budgets, bid strategies', 'Read only'),
    ('AdGroup', 'View ad group structure and performance', 'Read only'),
    ('AdGroupAd', 'Review ad copy and ad-level performance metrics', 'Read only'),
    ('AdGroupCriterion', 'Analyze keyword performance, match types, Quality Score', 'Read only'),
    ('SearchTermView', 'Review actual search queries that triggered ads', 'Read only'),
    ('GeoTargetConstant', 'Understand geographic targeting settings', 'Read only'),
]
for i, (resource, purpose, ops) in enumerate(api_rows, 1):
    api_table.rows[i].cells[0].text = resource
    api_table.rows[i].cells[1].text = purpose
    api_table.rows[i].cells[2].text = ops

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run(
    'We do NOT modify, create, or delete any campaigns, ads, or keywords via the API. '
)
run.bold = True
p.add_run(
    'All campaign management is done through the Google Ads web interface. '
    'The API is used exclusively for reporting and analysis.'
)

# -- Example Queries --
doc.add_heading('Example GAQL Queries', level=2)

queries = [
    ('Campaign Performance Report',
     'SELECT\n'
     '  campaign.name,\n'
     '  campaign.status,\n'
     '  campaign_budget.amount_micros,\n'
     '  metrics.impressions,\n'
     '  metrics.clicks,\n'
     '  metrics.cost_micros,\n'
     '  metrics.conversions,\n'
     '  metrics.average_cpc,\n'
     '  metrics.ctr\n'
     'FROM campaign\n'
     'WHERE segments.date DURING LAST_30_DAYS\n'
     'ORDER BY metrics.cost_micros DESC'),
    ('Keyword Performance Report',
     'SELECT\n'
     '  ad_group_criterion.keyword.text,\n'
     '  ad_group_criterion.keyword.match_type,\n'
     '  metrics.impressions,\n'
     '  metrics.clicks,\n'
     '  metrics.cost_micros,\n'
     '  metrics.conversions,\n'
     '  metrics.average_cpc\n'
     'FROM keyword_view\n'
     'WHERE segments.date DURING LAST_30_DAYS\n'
     '  AND metrics.impressions > 0\n'
     'ORDER BY metrics.cost_micros DESC'),
    ('Search Terms Report',
     'SELECT\n'
     '  search_term_view.search_term,\n'
     '  metrics.impressions,\n'
     '  metrics.clicks,\n'
     '  metrics.cost_micros,\n'
     '  metrics.conversions\n'
     'FROM search_term_view\n'
     'WHERE segments.date DURING LAST_30_DAYS\n'
     'ORDER BY metrics.impressions DESC'),
]

for query_title, query_text in queries:
    p = doc.add_paragraph()
    run = p.add_run(query_title + ':')
    run.bold = True

    code_p = doc.add_paragraph()
    code_run = code_p.add_run(query_text)
    code_run.font.name = 'Consolas'
    code_run.font.size = Pt(9)
    code_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    doc.add_paragraph('')  # spacer

# -- Tool Mockup --
doc.add_heading('Tool Mockup', level=1)
doc.add_paragraph(
    'The dashboard is an internal-only tool served on localhost. '
    'Below is a mockup of the reporting interface:'
)

# Embed the mockup image
mockup_path = os.path.join(os.path.dirname(__file__), 'dashboard-mockup.png')
if os.path.exists(mockup_path):
    doc.add_picture(mockup_path, width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
else:
    doc.add_paragraph('[Dashboard mockup image — see dashboard-mockup.png]')

doc.add_paragraph('')

features = doc.add_paragraph()
run = features.add_run('Dashboard Features:')
run.bold = True

feature_items = [
    'Account Overview — Summary cards showing total spend, clicks, impressions, '
    'conversions, cost/conversion, and CTR for the selected date range.',
    'Campaign Performance — Table comparing all campaigns with status indicators, '
    'click/impression/cost/conversion metrics.',
    'Top Keywords — Table showing keyword-level performance including match type, '
    'CTR, average CPC, and total cost.',
    'Date Range Selector — Dropdown to switch between Last 7 Days, Last 30 Days, '
    'Last 90 Days, and All Time views.',
]
for item in feature_items:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_paragraph()
run = p.add_run(
    'The tool is not externally accessible. It runs locally on the business owner\'s '
    'machine and is used solely for internal reporting and analysis.'
)
run.italic = True

# -- Rate Limiting --
doc.add_heading('Rate Limiting and Responsible Use', level=1)
rate_items = [
    'API calls are made periodically (every 4 hours during business hours), not in real-time.',
    'Results are cached locally to minimize redundant API requests.',
    'We anticipate fewer than 50 API requests per day.',
    'We will respect all rate limits and quotas defined by the Google Ads API Terms of Service.',
]
for item in rate_items:
    doc.add_paragraph(item, style='List Bullet')

# -- Contact --
doc.add_heading('Contact Information', level=1)
contact_table = doc.add_table(rows=4, cols=2, style='Light Grid Accent 1')
contact_table.alignment = WD_TABLE_ALIGNMENT.CENTER
contacts = [
    ('Business Owner', 'Jake Wells'),
    ('Email', 'jake@pianotuningkc.com'),
    ('Phone', '(913) 725-8124'),
    ('Website', 'https://pianotuningkc.com'),
]
for i, (label, value) in enumerate(contacts):
    contact_table.rows[i].cells[0].text = label
    contact_table.rows[i].cells[1].text = value

# Save
output_path = os.path.join(os.path.dirname(__file__), 'Google-Ads-API-Design-Document.docx')
doc.save(output_path)
print(f'Saved to: {output_path}')
